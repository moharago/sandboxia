"""문서 생성 서비스

draft 데이터를 기반으로 DOCX 문서를 생성합니다.
"""

import copy
import logging
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.table import Table
from docxtpl import DocxTemplate
from jinja2 import Environment, Undefined

from app.services.utils import is_flat_structure, unflatten


class SilentUndefined(Undefined):
    """undefined 변수를 빈 문자열로 조용히 처리

    템플릿에서 정의되지 않은 변수에 접근해도 에러 없이 빈 문자열 반환
    """
    def _fail_with_undefined_error(self, *args, **kwargs):
        return ""

    def __str__(self):
        return ""

    def __repr__(self):
        return ""

    def __html__(self):
        return ""

    def __getattr__(self, name):
        return SilentUndefined()

    def __getitem__(self, name):
        return SilentUndefined()


# docxtpl 렌더링용 Jinja2 환경 (undefined 에러 방지)
JINJA_ENV = Environment(undefined=SilentUndefined)

logger = logging.getLogger(__name__)


class SafeDict(dict):
    """없는 키 접근 시 빈 SafeDict 또는 빈 문자열 반환"""

    def __missing__(self, key):
        return SafeDict()

    def __str__(self):
        return ""

    def __repr__(self):
        if not self:
            return ""
        return super().__repr__()


class MarkerPreserver:
    """배열 마커 보존용 클래스

    docxtpl 렌더링 시 {{ prefix.field }} 형식의 마커를 그대로 유지합니다.
    렌더링 후 python-docx로 실제 값을 치환합니다.
    """
    def __init__(self, prefix: str):
        self.prefix = prefix

    def __getattr__(self, name: str) -> str:
        # Jinja2가 속성 접근 시 마커 문자열 반환
        return f"{{{{ {self.prefix}.{name} }}}}"

    def get(self, name: str, default: str = "") -> str:
        return f"{{{{ {self.prefix}.{name} }}}}"


def _to_safe_dict(data):
    """일반 dict를 SafeDict로 재귀 변환"""
    if isinstance(data, dict):
        return SafeDict({k: _to_safe_dict(v) for k, v in data.items()})
    if isinstance(data, list):
        return [_to_safe_dict(v) for v in data]
    return data

# 템플릿 디렉토리
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

# 트랙별 템플릿 파일 매핑
TEMPLATE_FILES = {
    "fastcheck": {
        "fastcheck-1": "fastcheck-1.docx",
        "fastcheck-2": "fastcheck-2.docx",
    },
    "temporary": {
        "temporary-1": "temporary-1.docx",
        "temporary-2": "temporary-2.docx",
        "temporary-3": "temporary-3.docx",
        "temporary-4": "temporary-4.docx",
    },
    "demonstration": {
        "demonstration-1": "demonstration-1.docx",
        "demonstration-2": "demonstration-2.docx",
        "demonstration-3": "demonstration-3.docx",
        "demonstration-4": "demonstration-4.docx",
    },
}


def _copy_table_row(table: Table, row_idx: int) -> None:
    """테이블의 특정 행을 복사하여 바로 아래에 추가"""
    row = table.rows[row_idx]
    # XML 요소 복사
    tr = row._tr
    new_tr = copy.deepcopy(tr)
    tr.addnext(new_tr)


def _find_marker_in_table(table: Table, marker: str) -> int | None:
    """테이블에서 마커 텍스트가 포함된 행의 인덱스 반환"""
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            if marker in cell.text:
                return idx
    return None


def _replace_text_in_row(row, old_text: str, new_text: str) -> None:
    """테이블 행의 모든 셀에서 텍스트 치환

    Word에서 텍스트가 여러 run으로 분리될 수 있으므로
    paragraph 전체 텍스트를 합쳐서 치환합니다.
    """
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            # paragraph 전체 텍스트에서 치환
            full_text = paragraph.text
            if old_text in full_text:
                # 모든 run의 텍스트를 합쳐서 치환
                new_full_text = full_text.replace(old_text, new_text)
                # 첫 번째 run에 새 텍스트, 나머지 run은 빈 문자열
                if paragraph.runs:
                    paragraph.runs[0].text = new_full_text
                    for run in paragraph.runs[1:]:
                        run.text = ""


def _expand_array_rows(doc: Document, array_data: list, field_prefix: str, fields: list[str]) -> None:
    """배열 데이터를 기반으로 테이블 행 확장

    템플릿에는 {field_prefix}0.{field} 형식의 플레이스홀더가 있는 한 행이 있음.
    이 행을 배열 길이만큼 복제하고 각 행에 데이터를 채움.
    """
    if not array_data:
        return

    # 마커 찾기 (예: "org0." 또는 "person0.")
    marker = f"{field_prefix}0."

    # 모든 테이블 검색 (중첩 테이블 포함)
    for table in doc.tables:
        _process_table_for_array(table, array_data, marker, field_prefix, fields)


def _process_table_for_array(table: Table, array_data: list, marker: str, field_prefix: str, fields: list[str]) -> None:
    """단일 테이블에서 배열 행 확장 처리"""
    # 중첩 테이블도 처리
    for row in table.rows:
        for cell in row.cells:
            for nested_table in cell.tables:
                _process_table_for_array(nested_table, array_data, marker, field_prefix, fields)

    # 마커가 있는 행 찾기
    row_idx = _find_marker_in_table(table, marker)
    if row_idx is None:
        return

    # 1단계: 먼저 행을 필요한 만큼 복제 (마커 치환 전에!)
    # 원본 행을 복제해야 마커가 보존됨
    for _ in range(1, len(array_data)):
        _copy_table_row(table, row_idx)  # 항상 원본 행(row_idx)을 복제

    # 2단계: 각 행의 마커를 해당 항목 데이터로 치환
    for i, item in enumerate(array_data):
        current_row = table.rows[row_idx + i]
        for field in fields:
            old_placeholder = f"{{{{ {field_prefix}0.{field} }}}}"
            new_value = str(item.get(field, ""))
            _replace_text_in_row(current_row, old_placeholder, new_value)


def generate_docx(
    track: str,
    form_id: str,
    draft_data: dict,
) -> BytesIO:
    """draft 데이터로 DOCX 문서 생성

    Args:
        track: 트랙 유형 (fastcheck, temporary, demonstration)
        form_id: 폼 ID (fastcheck-1, temporary-1 등)
        draft_data: 해당 폼의 draft 데이터

    Returns:
        BytesIO: 생성된 DOCX 파일 (메모리)
    """
    # 템플릿 파일 경로
    template_files = TEMPLATE_FILES.get(track, {})
    template_name = template_files.get(form_id)

    if not template_name:
        raise ValueError(f"템플릿이 없습니다: {track}/{form_id}")

    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"템플릿 파일이 없습니다: {template_path}")

    # 템플릿 로드
    doc = DocxTemplate(template_path)

    # draft 데이터를 템플릿 컨텍스트로 변환
    context = _build_context(draft_data)

    logger.info("문서 생성: track=%s, form_id=%s", track, form_id)

    # 렌더링 (undefined 변수는 빈 문자열로 처리)
    doc.render(context, jinja_env=JINJA_ENV)

    # demonstration-2, temporary-2의 경우 배열 행 확장
    if form_id in ("demonstration-2", "temporary-2"):
        # python-docx Document로 변환
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        pydoc = Document(buffer)

        # 원본 draft_data에서 배열 추출 (context가 아닌 원본 사용)
        if is_flat_structure(draft_data):
            draft_data = unflatten(draft_data)

        # applicantOrganizations 확장
        orgs = draft_data.get("applicantOrganizations", [])
        if orgs:
            _expand_array_rows(
                pydoc, orgs, "org",
                ["organizationName", "organizationType", "responsiblePersonName",
                 "position", "phoneNumber", "email"]
            )

        # signers (submission) 확장
        signers = draft_data.get("submission", [])
        if signers:
            _expand_array_rows(
                pydoc, signers, "signer",
                ["organizationName", "name", "signature"]
            )

        # keyPersonnel 확장
        # 템플릿 플레이스홀더: {{ person0.qualifications }}, {{ person0.experience }}
        personnel = draft_data.get("keyPersonnel", [])
        if personnel:
            # 데이터 필드명을 템플릿 플레이스홀더에 맞게 매핑
            mapped_personnel = []
            for p in personnel:
                if not isinstance(p, dict):
                    continue
                mapped = {
                    "name": p.get("name", ""),
                    "department": p.get("department", ""),
                    "position": p.get("position", ""),
                    "responsibilities": p.get("responsibilities", ""),
                    # 템플릿: qualifications (데이터: qualifications 또는 qualificationsOrSkills)
                    "qualifications": p.get("qualifications") or p.get("qualificationsOrSkills", ""),
                    # 템플릿: experienceYears (데이터: experience 또는 experienceYears)
                    "experienceYears": p.get("experienceYears") or p.get("experience", ""),
                }
                mapped_personnel.append(mapped)

            _expand_array_rows(
                pydoc, mapped_personnel, "person",
                ["name", "department", "position", "responsibilities",
                 "qualifications", "experienceYears"]
            )

        # BytesIO로 저장
        buffer = BytesIO()
        pydoc.save(buffer)
        buffer.seek(0)
        return buffer

    # BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def _replace_none(data):
    """None 값을 빈 문자열로 재귀 변환"""
    if isinstance(data, dict):
        return {k: _replace_none(v) for k, v in data.items()}
    if isinstance(data, list):
        return [_replace_none(v) for v in data]
    return "" if data is None else data


def _format_date_korean(date_str: str) -> str:
    """날짜를 한국어 형식으로 변환

    지원 형식:
    - 2026-02-02 → 2026년 2월 2일
    - 2026. 02. 11. → 2026년 2월 11일
    - 2026년 2월 2일 → 그대로 반환
    """
    if not date_str:
        return ""
    try:
        date_str = str(date_str).strip()

        # 이미 한국어 형식이면 그대로 반환
        if "년" in date_str and "월" in date_str:
            return date_str

        # ISO 형식 (2026-02-02)
        if "-" in date_str:
            parts = date_str.split("-")
            if len(parts) == 3:
                year, month, day = parts
                return f"{year}년 {int(month)}월 {int(day)}일"

        # 점 형식 (2026. 02. 11. 또는 2026.02.11)
        if "." in date_str:
            # 공백과 점 제거 후 파싱
            cleaned = date_str.replace(" ", "").rstrip(".")
            parts = cleaned.split(".")
            if len(parts) == 3:
                year, month, day = parts
                return f"{year}년 {int(month)}월 {int(day)}일"

        return date_str
    except Exception:
        return str(date_str)


def _build_context(draft_data: dict) -> dict:
    """draft 데이터를 템플릿 컨텍스트로 변환"""
    # flat 구조면 nested로 변환
    if is_flat_structure(draft_data):
        draft_data = unflatten(draft_data)

    # None → 빈 문자열, dict → SafeDict
    context = _to_safe_dict(_replace_none(draft_data))

    # 날짜 필드 한국어 형식 변환
    if "application" in context:
        app_date = context["application"].get("applicationDate", "")
        context["application"]["applicationDate"] = _format_date_korean(app_date)

    if "submissionDate" in context:
        sub_date = context["submissionDate"].get("submissionDate", "")
        context["submissionDate"]["submissionDate"] = _format_date_korean(sub_date)

    if "projectInfo" in context:
        period = context["projectInfo"].get("period", {})
        if isinstance(period, dict):
            if "startDate" in period:
                period["startDate"] = _format_date_korean(period["startDate"])
            if "endDate" in period:
                period["endDate"] = _format_date_korean(period["endDate"])

    if "organizationProfile" in context:
        general_info = context["organizationProfile"].get("generalInfo", {})
        if isinstance(general_info, dict) and "establishmentDate" in general_info:
            general_info["establishmentDate"] = _format_date_korean(general_info["establishmentDate"])

    # 체크박스 필드 (√ 또는 빈값)
    service_type = draft_data.get("technologyService", {}).get("type", "")
    context["checkTechnology"] = "√" if service_type == "technology" else ""
    context["checkService"] = "√" if service_type == "service" else ""
    context["checkTechnologyAndService"] = "√" if service_type == "technologyAndService" else ""

    # 중첩 구조 평탄화 (section.field → field)
    for key in ["technologyServiceDetails", "legalIssues", "additionalQuestions"]:
        if isinstance(context.get(key), dict):
            context[key] = context[key].get(key, "")

    # 임시허가 신청 사유 체크박스 (temporary-1)
    permit_reasons = draft_data.get("temporaryPermitReason", {}).get("temporaryPermitReason", [])
    context["checkNoApplicableStandards"] = "√" if "noApplicableStandards" in permit_reasons else ""
    context["checkUnclearStandards"] = "√" if "unclearOrUnreasonableStandards" in permit_reasons else ""

    # 임시허가 신청 사유 해당여부 (temporary-3)
    eligibility = draft_data.get("eligibility", {})
    context["checkEligibilityNoStandards"] = "O" if eligibility.get("noApplicableStandards") else ""
    context["checkEligibilityUnclear"] = "O" if eligibility.get("unclearOrUnreasonableStandards") else ""

    # 실증특례 신청 사유 체크박스 (demonstration-1)
    reg_exemption = draft_data.get("regulatoryExemptionReason", {})
    reason1 = reg_exemption.get("reason1_impossibleToApplyPermit", False)
    reason2 = reg_exemption.get("reason2_unclearOrUnreasonableCriteria", False)
    context["checkReason1"] = "√" if reason1 else ""
    context["checkReason2"] = "√" if reason2 else ""

    # 실증특례 소명서 체크박스 (demonstration-3)
    demo_eligibility = draft_data.get("eligibility", {})
    context["checkEligibilityImpossible"] = "O" if demo_eligibility.get("impossibleToApplyPermitByOtherLaw") else ""
    context["checkEligibilityUnclear"] = "O" if demo_eligibility.get("unclearOrUnreasonableCriteria") else ""

    # 템플릿에서 직접 비교하는 경우를 위한 변환 (True/False → 체크 표시)
    # {{ regulatoryExemptionReason.reason1_impossibleToApplyPermit }} 등
    if "regulatoryExemptionReason" in context:
        context["regulatoryExemptionReason"]["reason1_impossibleToApplyPermit"] = "√" if reason1 else ""
        context["regulatoryExemptionReason"]["reason2_unclearOrUnreasonableCriteria"] = "√" if reason2 else ""

    # {{ technologyService.type == "xxx" }} 직접 비교 대응
    if "technologyService" in context:
        ts = context["technologyService"]
        ts_type = ts.get("type", "")
        # jinja2 조건식 결과가 True/False로 나오는 것 방지
        # 템플릿에서 직접 비교하면 True/False가 출력되므로 미리 변환
        context["technologyService"]["type_is_technology"] = "√" if ts_type == "technology" else ""
        context["technologyService"]["type_is_service"] = "√" if ts_type == "service" else ""
        context["technologyService"]["type_is_both"] = "√" if ts_type == "technologyAndService" else ""

    # 반복 테이블용 빈 배열 보장
    array_fields = [
        "applicantOrganizations",  # demonstration-2, temporary-2 신청기관
        "submission",              # demonstration-2, temporary-2 서명
        "keyPersonnel",            # demonstration-2, temporary-2 주요인력
    ]
    for field in array_fields:
        if field not in context or not isinstance(context.get(field), list):
            context[field] = []

    # 템플릿-JSON 필드명 불일치 alias
    context["signers"] = context.get("submission", [])

    # 배열 마커 보존 - docxtpl 렌더링 후에도 {{ xxx0.field }} 형식 유지
    # 실제 값 치환은 python-docx로 _expand_array_rows에서 수행
    # MarkerPreserver는 {{ prefix.field }} 형식의 마커 문자열을 반환
    context["org0"] = MarkerPreserver("org0")
    context["signer0"] = MarkerPreserver("signer0")
    context["person0"] = MarkerPreserver("person0")

    # organizationProfile.keyPersonnel → 최상위 keyPersonnel 연결
    if "keyPersonnel" in context and context.get("keyPersonnel"):
        if "organizationProfile" not in context:
            context["organizationProfile"] = SafeDict()
        context["organizationProfile"]["keyPersonnel"] = context.get("keyPersonnel", [])

    # financialStatus → organizationProfile.financialStatus 연결 (템플릿 호환)
    # yearM1 → year1, yearM2 → year2 키 변환
    if "financialStatus" in context and isinstance(context["financialStatus"], dict):
        converted_financial = {}
        for field_name, field_data in context["financialStatus"].items():
            if isinstance(field_data, dict):
                converted_financial[field_name] = {
                    "year1": field_data.get("yearM1", field_data.get("year1", "")),
                    "year2": field_data.get("yearM2", field_data.get("year2", "")),
                    "average": field_data.get("average", ""),
                }
            else:
                converted_financial[field_name] = field_data

        if "organizationProfile" not in context:
            context["organizationProfile"] = SafeDict()
        context["organizationProfile"]["financialStatus"] = converted_financial

    # humanResources → organizationProfile.humanResources 연결 (템플릿 호환)
    if "humanResources" in context and isinstance(context["humanResources"], dict):
        if "organizationProfile" not in context:
            context["organizationProfile"] = SafeDict()
        context["organizationProfile"]["humanResources"] = context["humanResources"]

    # 모든 템플릿에서 사용하는 중첩 필드 기본값 보장
    default_nested_fields = [
        # 공통 - 신청자 정보
        "applicant",
        "application",
        # demonstration-2, temporary-2
        "projectInfo",
        "technologyService",
        "regulatoryExemption",
        "testPlan",
        "operationPlan",
        "expectedEffects",
        "postTestPlan",
        "organizationAndBudget",
        "organizationProfile",
        "financialStatus",
        "humanResources",
        "submissionDate",
        "temporaryPermitRequest",
        "businessPlan",
        "expansionPlan",
        # demonstration-3, temporary-3
        "eligibility",
        "justification",
        # demonstration-4, temporary-4
        "protectionAndResponse",
        "riskAndMitigation",
        "stakeholderConflict",
        "safetyVerification",
        "userProtectionPlan",
        "riskAndResponse",
        "stakeholderConflictResolution",
    ]
    for field in default_nested_fields:
        if field not in context:
            context[field] = SafeDict()

    return context
